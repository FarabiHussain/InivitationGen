import os
import datetime
import variables as vars
from docx import Document
from docx.shared import Cm as CM
from docx.shared import Pt as PT
from CTkMessagebox import CTkMessagebox as popup
from path_manager import resource_path
from utils import *
from icecream import ic
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Entity import *


# initializes the fill info, output and input files
def process_data():
    date_on_document = datetime.datetime.now()
    initialize_entities()

    form_data = {
        '[DAY]': format_day(date_on_document.strftime("%d")),
        '[MONTH]': date_on_document.strftime("%B"),
        '[YEAR]': date_on_document.strftime("%Y"),
    }

    # set the input and output files
    input_file = resource_path("assets\\templates\\double.docx")
    output_file = f"{(datetime.datetime.now().timestamp())}.docx"

    form_data['[GUEST_INFO]'] = ""


    for i in range(len(hosts_list)):
        host = hosts_list[i]

        form_data[f'[HOST{i+1}_NAME]'] = host.get('name')
        form_data[f'[HOST{i+1}_BIRTH]'] = host.get('birth')
        form_data[f'[HOST{i+1}_STATUS]'] = host.get('status')
        form_data[f'[HOST{i+1}_PASSPORT]'] = host.get('passport')
        form_data[f'[HOST{i+1}_ADDRESS]'] = host.get('address')
        form_data[f'[HOST{i+1}_PHONE]'] = host.get('phone')
        form_data[f'[HOST{i+1}_OCCUPATION]'] = host.get('occupation')
        form_data[f'[HOST{i+1}_EMAIL]'] = host.get('email')
        form_data[f'[HOST{i+1}_RELATION_TO_HOST{"2" if (i+1)==1 else "1"}]'] = host.get('relation_to_other_host').lower()

        if (i==0):
            form_data['[BEARER]'] = hosts_list[0].get('bearer')

    for i in range(len(guests_list)):
        guest = guests_list[i]

        form_data['[GUEST_NAMES]'] = format_collective_names(guests_list)
        form_data['[GUEST_INFO]'] += f"{"\n" if i > 0 else ""}{guest.get('name')}, born {guest.get('birth')}, Passport Number {guest.get('passport')}"
        form_data[f'[GUEST{i+1}_BIRTH]'] = guest.get('birth')
        form_data[f'[GUEST{i+1}_CITIZENSHIP]'] = guest.get('citizenship')
        form_data[f'[GUEST{i+1}_PASSPORT]'] = guest.get('passport')
        form_data[f'[GUEST{i+1}_ADDRESS]'] = guest.get('address')
        form_data[f'[GUEST{i+1}_PHONE]'] = guest.get('phone')
        form_data[f'[GUEST{i+1}_OCCUPATION]'] = guest.get('occupation')
        form_data[f'[GUEST{i+1}_RELATIONSHIP]'] = guest.get('relationship_to_host1')
        form_data[f'[GUEST{i+1}_CANADIAN_ADDRESS]'] = guest.get('canadian_address')

        if i==0:
            form_data[f'[GUEST_PURPOSE]'] = guests_list[0].get('purpose')
            form_data[f'[GUEST_ARRIVAL]'] = guests_list[0].get('arrival')
            form_data[f'[GUEST_DEPARTURE]'] = guests_list[0].get('departure')

    return {
        'form_data': form_data, 
        'input_file': input_file, 
        'output_file': output_file
    }


# create objects of hosts and guests
def initialize_entities():
    global guests_list, hosts_list

    guests_list = []
    hosts_list = []
    guest_fields = vars.field_dicts[0]
    host_fields = vars.field_dicts[1]

    # append guests
    for i in range(1,4):
        guests_list.append(Guest({
            'name': guest_fields[f'guest{i}_entry_name'].get(),
            'birth': guest_fields[f'guest{i}_datepicker_birth'].get(),
            'address': guest_fields[f'guest{i}_entry_address'].get(),
            'phone': guest_fields[f'guest{i}_entry_phone'].get(),
            'occupation': guest_fields[f'guest{i}_entry_occupation'].get(),
            'relationship_to_host1': guest_fields[f'guest{i}_entry_relationship_to_host1'].get(),
            'purpose': guest_fields[f'guest{i}_entry_purpose'].get(),
            'canadian_address': guest_fields[f'guest{i}_entry_canadian_address'].get(),
            'passport': guest_fields[f'guest{i}_entry_passport_number'].get(),
            'citizenship': guest_fields[f'guest{i}_entry_citizenship'].get(),

            # the arrival and departure dates are always the same as guest1
            'arrival': guest_fields[f'guest1_datepicker_arrival'].get(),
            'departure': guest_fields[f'guest1_datepicker_departure'].get(),
        }))

    # append hosts
    for i in range(1,3):
        hosts_list.append(Host({
            'name': host_fields[f'host{i}_entry_name'].get(),
            'birth': host_fields[f'host{i}_datepicker_birth'].get(),
            'status': host_fields[f'host{i}_entry_status'].get(),
            'passport': host_fields[f'host{i}_entry_passport_number'].get(),
            'address': host_fields[f'host{i}_entry_address'].get(),
            'phone': host_fields[f'host{i}_entry_phone'].get(),
            'occupation': host_fields[f'host{i}_entry_occupation'].get(),
            'email': host_fields[f'host{i}_entry_email'].get(),
            'relation_to_other_host': host_fields[f'host{"1" if i == 1 else "2"}_entry_relationship_to_host{"2" if i == 1 else "1"}'].get(),

            # the bearer of expenses and attached documents are always the same as host1 
            'bearer': host_fields[f'host1_combo_bearer_of_expenses'].get(),
            'attached': host_fields[f'host1_entry_attached_documents'].get(),
        }))


# generate the docx with the input info
def generate_doc():
    data = None
    doc = None
    guest_fields = vars.field_dicts[0]
    host_fields = vars.field_dicts[1]
    # finances_fields = vars.field_dicts[2]

    # initiate the data and document
    try:
        data = process_data()
        doc = Document(data['input_file'])
    except Exception as e:
        ErrorPopup(msg=f'Exception in process_data(): {str(e)}')
        return False

    # replace placeholders on the document
    try:
        current_replace = ""
        for paragraph in doc.paragraphs:
            for key, value in data['form_data'].items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        current_replace = value
                        run.text = run.text.replace(key, value)
    except Exception as e:
        ErrorPopup(msg=f'Exception while editing document\n\nPlacing {current_replace}\n\n{e}')
        return False

    # invitation letter intro
    insert_paragraph(
        paragraph=doc.add_paragraph(),
        text=(
            f"This letter is to express me and my {host_fields['host2_entry_relationship_to_host1'].get()}'s interest in inviting {guest_fields['guest1_entry_name'].get()} Canada " + 
            f"and to furthermore support the Temporary Resident Visa application."
        )
    )

    # table containing hosts' details
    for index in range(1,3):
        insert_table(
            document=doc, 
            table_heading="\nMy details are as follows," if index==1 else f"\n\nMy {host_fields['host1_entry_relationship_to_host2'].get()}'s details are as follows,", 
            table_items=[
                {"label": "Full Name", "info": host_fields[f'host{index}_entry_name'].get()},
                {"label": "Date of Birth", "info": host_fields[f'host{index}_datepicker_birth'].get()},
                {"label": "Canadian Status", "info": host_fields[f'host{index}_entry_status'].get()},
                {"label": "Passport Number", "info": host_fields[f'host{index}_entry_passport_number'].get()},
                {"label": "Residential Address", "info": host_fields[f'host{index}_entry_address'].get()},
                {"label": "Phone Number", "info": host_fields[f'host{index}_entry_phone'].get()},
                {"label": "Current Occupation", "info": host_fields[f'host{index}_entry_occupation'].get()},
                {"label": "Email Address", "info": host_fields[f'host{index}_entry_email'].get()},
            ]
        )

    # table containing guests' details
    for i in range(len(guests_list)):
        
        if i==0:
            doc.add_page_break()

        guest = guests_list[i]

        # add the table containing guest's details
        insert_table(
            document=doc, 
            table_heading=f"The details of {guest.get('name')} are as follows,", 
            table_items=[
                {"label": "Full Name", "info": guest.get('name')},
                {"label": "Date of Birth", "info": guest.get('birth')},
                {"label": "Citizen of", "info": guest.get('citizenship')},
                {"label": "Passport No.", "info": guest.get('passport')},
                {"label": "Residential Address", "info": guest.get('address')},
                {"label": "Phone Number", "info": guest.get('phone')},
                {"label": "Current Occupation", "info": guest.get('occupation')},
                {"label": "Relationship to Inviter", "info": guest.get('relationship_to_host1')},
                {"label": "Purpose of Visit", "info": guest.get('purpose')},
                {"label": "Arrival Date", "info": guest.get('arrival')},
                {"label": "Departure Date", "info": guest.get('departure')},
                {"label": "Primary Residence in Canada", "info": guest.get('canadian_address')},
            ]
        )

        if i < 3:
            insert_paragraph(
                paragraph=doc.add_paragraph(),
                text=("")
            )

    bearer = host_fields["host1_combo_bearer_of_expenses"].get()
    attached = host_fields["host1_entry_attached_documents"].get()

    outro_text=(
        f"The airfare, travel expenses, would be borne by {format_collective_names(guests_list)}. " +
        f"All expenses in connection with their visit to Canada will be {bearer}. "
    )

    # attached documents
    if len(attached) > 0:
        outro_text += f"\n\nAttached with the application are {attached}. "

    outro_text += f"\n\nIf any clarification or information is required, please do not hesitate to contact us at our email addresses and phone numbers below."

    insert_paragraph(paragraph=doc.add_paragraph(), text=(outro_text))

    # table for hosts to sign
    insert_table(
        document=doc, 
        table_heading="\n\n\n\n",
        table_items=[
            {"label": "_________________________________", "info": "_________________________________"},
            {"label": host_fields['host1_entry_name'].get(), "info": host_fields['host2_entry_name'].get()},
            {"label": host_fields['host1_entry_email'].get(), "info": host_fields['host2_entry_email'].get()},
            {"label": host_fields['host1_entry_phone'].get(), "info": host_fields['host2_entry_phone'].get()},
        ],
        table_props={
            'color': '#ffffff',
            "cell_alignment": [
                WD_ALIGN_PARAGRAPH.LEFT, 
                WD_ALIGN_PARAGRAPH.LEFT
            ]
        }
    )

    if (user_confirmation()):
        # create the output file
        save_doc(doc, data)


# prompts user for confirmation before generating the document
def user_confirmation():
    entity_types = [
        "guest 1",
        "guest 2",
        "guest 3",
        "host 1",
        "host 2",
    ]

    for i, entity in enumerate(guests_list + hosts_list):
        response = entity.is_filled()

        if response[0] is not True:

            # print(list((set(entity.get_valid_props())) - set(response[1])))

            # if list((set(entity.get_valid_props())) - set(response[1])) == ["arrival", "departure", "birth"]:
            #     msg = f"{entity_types[i].capitalize()} is empty.\n\nContinue?"
            # else:
            #    msg = f"The following field(s) in {entity_types[i].capitalize()} are not filled:\n\n     • {("\n     • ").join(response[1])}\n\nContinue?"

            msg = f"The following field(s) in {entity_types[i].capitalize()} are not filled:\n\n     • {("\n     • ").join(response[1])}\n\nContinue?"

            prompt = PromptPopup(msg=msg)

            if prompt.get() is not True:
                # InfoPopup(msg="Operation cancelled.")
                return False

    return True


# set up folders and save files, print if needed
def save_doc(doc, data):
    try:
        # set up the output directory
        output_dir = os.getcwd() + "\\output\\"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # save the file to the output folder
        doc.save(output_dir + data['output_file'])

        # open the word file
        os.startfile(output_dir + data['output_file'])

        # return the filename
        return data['output_file']

    except Exception as e:
        ErrorPopup(msg=f"Exception when saving and opening document:\n\n{str(e)}")
        return False


# insert table with the passed data
def insert_table(document, table_heading=None, table_items=[], table_props=None):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    if table_props is None:
        table_props = {
            "color": "#AAAAAA",
            "cell_alignment": [
                WD_ALIGN_PARAGRAPH.LEFT, 
                WD_ALIGN_PARAGRAPH.RIGHT
            ]
        }

    # add the table heading
    try:
        if table_heading is not None:
            insert_paragraph(
                paragraph=document.add_paragraph(),
                text=table_heading,
                bolded=True,
                italicized=False
            )
    except Exception as e:
        popup(title="", message=f'Exception when adding table heading: {table_heading}\n\n{e}', corner_radius=4)
        return False

    # add the table contents
    try:
        # Creating a table object
        table_obj = document.add_table(rows=len(table_items), cols=len(table_items[0]))

        # add rows
        for idx, row in enumerate(table_obj.rows):
            curr_row = row.cells
            row.height = CM(0.50)
            curr_row[0].text = table_items[idx]['label']
            curr_row[0].paragraphs[0].alignment = table_props['cell_alignment'][0]
            curr_row[1].text = table_items[idx]['info']
            curr_row[1].paragraphs[0].alignment = table_props['cell_alignment'][1]

        # set column widths and borders
        for idx, _ in enumerate(table_items[0]):
            for cell in (table_obj.columns[idx].cells):
                cell.width = CM(20)
                set_cell_border(cell, top={"sz": 1, "color": table_props['color'], "val": "single", "space": "4"})
                set_cell_border(cell, bottom={"sz": 1, "color": table_props['color'], "val": "single", "space": "0"})
    except Exception as e:
        popup(title="", message=f'Exception when adding table contents\n\n{e}', corner_radius=4)
        return False


# insert a new paragraph
def insert_paragraph(paragraph, text="", bolded=False, italicized=False):

    if bolded or italicized:
        runner = paragraph.add_run(text)
        runner.bold = bolded
        runner.italic = italicized

    else:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        paragraph.paragraph_format.space_before = PT(0)
        paragraph.paragraph_format.space_after = PT(0)
        new_para = Paragraph(new_p, paragraph._parent)
        new_para.add_run(text)
