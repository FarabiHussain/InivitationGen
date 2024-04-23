import datetime, os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from GUI import *
from docx.shared import Cm as CM
from docx.shared import Pt as PT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# format date to `{date + suffix} {full month name} {year}`
def format_date(date_string):
    if date_string == 'advance':
        return date_string

    temp = datetime.datetime.strptime(date_string, '%d/%m/%Y')
    return str(format_day(temp.strftime("%d")) + " " + temp.strftime("%B") + " " + temp.strftime("%Y"))


# add suffix 'th'/'st'/'rd' to days
def format_day(day):
    day = int(day)
    suffix = ""

    if (4 <= day <= 20 or 24 <= day <= 30):
        suffix = "th"
    else:
        suffix = ["st", "nd", "rd"][day % 10 - 1]

    return str(day) + suffix


# format phone number to add brackets and hyphens when applicable
def format_phone(number):
    if (len(number) == 10):  
        area = number[0] + number[1] + number[2]
        prefix = number[3] + number[4] + number[5]
        line = number[6] + number[7] + number[8] + number[9]
        return "(" + area + ")" + " " + prefix + "-" + line

    return number


# 
def format_collective_names(guests:dict = {}):
    temp = ""

    if guests['guest_1'] is not None:
        temp = f"{guests['guest_1'].get('name')}"
    if guests['guest_1'] is not None and guests['guest_2'] is not None:
        temp = f"{guests['guest_1'].get('name')} and {guests['guest_2'].get('name')}"
    if guests['guest_1'] is not None and guests['guest_2'] is not None and guests['guest_3'] is not None:
        temp = f"{guests['guest_1'].get('name')}, {guests['guest_2'].get('name')}, and {guests['guest_3'].get('name')}"

    return temp


# 
def set_cell_color(cell, color="ffffff"):
    cell_xml = cell._tc
    cell_props = cell_xml.get_or_add_tcPr()
    shade_obj = OxmlElement("w:shd")
    shade_obj.set(qn("w:fill"), color)
    cell_props.append(shade_obj)


# 
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


# 
def set_row_color(row, color="FFFFFF"):
    for cell in row.cells:
        set_cell_color(cell, color)


# 
def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


# set up folders and save files, print if needed
def save_doc(doc, data):
    try:
        # set up the output directory
        output_dir = os.getcwd() + "\\output\\"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # save the file to the output folder
        doc.save(output_dir + data['output_file'])

        # open the word file
        os.startfile(output_dir + data['output_file'])

        # return the filename
        return data['output_file']

    except Exception as e:
        ErrorPopup(msg=f"Exception when saving and opening document:\n\n{str(e)}")
        return False


# insert table with the passed data
def insert_table(document, table_heading=None, table_items=[], table_props=None):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    if table_props is None:
        table_props = {
            "color": "#AAAAAA",
            "cell_alignment": [
                WD_ALIGN_PARAGRAPH.LEFT, 
                WD_ALIGN_PARAGRAPH.RIGHT
            ]
        }

    # add the table heading
    try:
        if table_heading is not None:
            insert_paragraph(
                paragraph=document.add_paragraph(),
                text=table_heading,
                bolded=True,
                italicized=False
            )
    except Exception as e:
        ErrorPopup(f'Exception when adding table heading: {table_heading}\n\n{e}')
        return False

    # add the table contents
    try:
        # Creating a table object
        table_obj = document.add_table(rows=len(table_items), cols=len(table_items[0]))

        # add rows
        for idx, row in enumerate(table_obj.rows):
            curr_row = row.cells
            row.height = CM(0.50)
            curr_row[0].text = table_items[idx]['label']
            curr_row[0].paragraphs[0].alignment = table_props['cell_alignment'][0]
            curr_row[1].text = table_items[idx]['info']
            curr_row[1].paragraphs[0].alignment = table_props['cell_alignment'][1]

        # set column widths and borders
        for idx, _ in enumerate(table_items[0]):
            for cell in (table_obj.columns[idx].cells):
                cell.width = CM(20)
                set_cell_border(cell, top={"sz": 1, "color": table_props['color'], "val": "single", "space": "4"})
                set_cell_border(cell, bottom={"sz": 1, "color": table_props['color'], "val": "single", "space": "0"})
    except Exception as e:
        ErrorPopup(f'Exception when adding table contents\n\n{e}')
        return False


# insert a new paragraph
def insert_paragraph(paragraph, text="", bolded=False, italicized=False):

    if bolded or italicized:
        runner = paragraph.add_run(text)
        runner.bold = bolded
        runner.italic = italicized

    else:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        paragraph.paragraph_format.space_before = PT(0)
        paragraph.paragraph_format.space_after = PT(0)
        new_para = Paragraph(new_p, paragraph._parent)
        new_para.add_run(text)


